"""Convert technical-proposal.md to a formatted .docx file."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# Normal style
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Heading styles
for i, size in enumerate([22, 16, 14, 12], 1):
    h_style = doc.styles[f'Heading {i}']
    h_font = h_style.font
    h_font.name = 'Arial'
    h_font.size = Pt(size)
    h_font.bold = True
    h_font.color.rgb = RGBColor(0, 51, 102)
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

with open('docs/technical-proposal.md', 'r', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')
i = 0
in_code_block = False
code_lines = []

while i < len(lines):
    line = lines[i]

    # Code block start/end
    if line.startswith('```'):
        if in_code_block:
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(60, 60, 60)
            in_code_block = False
            code_lines = []
        else:
            in_code_block = True
        i += 1
        continue

    if in_code_block:
        code_lines.append(line)
        i += 1
        continue

    # Table rows
    if line.startswith('|') and line.strip().endswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].startswith('|') and lines[i].strip().endswith('|'):
            table_lines.append(lines[i])
            i += 1

        # Parse cells, skip separator rows like |---|---|
        data_rows = []
        for tl in table_lines:
            cells = [c.strip() for c in tl.split('|')[1:-1]]
            if all(re.match(r'^[-:]+$', c) for c in cells):
                continue
            data_rows.append(cells)

        if data_rows:
            num_cols = max(len(r) for r in data_rows)
            table = doc.add_table(rows=len(data_rows), cols=num_cols, style='Table Grid')
            table.autofit = True
            for ri, row_data in enumerate(data_rows):
                for ci in range(num_cols):
                    cell_text = row_data[ci] if ci < len(row_data) else ''
                    cell = table.rows[ri].cells[ci]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(2)
                        paragraph.paragraph_format.space_after = Pt(2)
                        for run in paragraph.runs:
                            run.font.size = Pt(9.5)
                            if ri == 0:
                                run.font.bold = True
            doc.add_paragraph()
        continue

    # Headings
    if line.startswith('# ') and not line.startswith('## '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## ') and not line.startswith('### '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### ') and not line.startswith('#### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('#### '):
        doc.add_heading(line[5:], level=4)
    elif line.strip() == '---':
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.add_run('─' * 50).font.color.rgb = RGBColor(200, 200, 200)
    elif line.strip():
        # Regular paragraph with inline formatting
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)

        text = line
        # Split by **bold** markers
        tokens = re.split(r'(\*\*[^*]+\*\*)', text)
        for token in tokens:
            if token.startswith('**') and token.endswith('**'):
                run = p.add_run(token[2:-2])
                run.font.bold = True
            else:
                # Split by `code` markers
                code_tokens = re.split(r'(`[^`]+`)', token)
                for ct in code_tokens:
                    if ct.startswith('`') and ct.endswith('`'):
                        run = p.add_run(ct[1:-1])
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(180, 50, 50)
                    else:
                        p.add_run(ct)

    i += 1

output_path = 'docs/联邦分布式查询系统-技术方案.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
